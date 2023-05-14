from tkinter import *
import requests
import json
from bs4 import BeautifulSoup
import docx
from PIL import ImageTk, Image
from tkinter import messagebox as msgbox

class wiki():
    def __init__(self, root):
        root.title("Find Any Person..")
        self.root = Frame(root,width=1000,height=600,bg="black")
        self.root.pack()
        self.main_win()

    # ============= Creating Main Window ===================
    def main_win(self):
        self.frame = Frame(self.root).pack()
        title = Label(self.frame, text="Get Details of Any Person", font=(
            "verdana", 18, "bold"), bg="lightblue", borderwidth=7, relief=SUNKEN, height=2)
        title.pack(side=TOP, fill=X)
        
        label = Label(self.frame, text="Enter The name Of Persone You Want to find : ", font=(
            "verdana 16 bold"), bg="black", fg="white")
        label.place(x=40, y=140)

        self.person_in = Entry(self.frame, width=35, font=(
            "verdana 16 bold"), borderwidth=5, relief=SUNKEN)
        self.person_in.place(x=90, y=200, height=50)

        search_btn = Button(self.frame, text="Search", font=(
            "verdana 16 bold"), borderwidth=5, relief=SUNKEN, command=self.search)
        search_btn.place(x=650, y=200, height=50)

    # =================== Collecting Data and Diplaying It =================
    def search(self):
        try:
            messagebox = Message(self.frame, text="Wait for While.. We Are Downloading the data...", font=(
                "Verdana", 12, "bold"), bg="black", fg="Red", width=650)
            messagebox.place(x=90, y=250)
            self.root.update()

            name = self.person_in.get()
            self.person_in.delete(0, END)
            self.paragraph = []

            url = f"https://en.wikipedia.org/wiki/{name}"
            result = requests.get(url)
            soup = BeautifulSoup(result.text, 'html.parser')

            # ============ getting paragraphs and Full Name ================
            par = soup.find_all('p')
            for i in par:
                if len(i.text) > 50:
                    self.paragraph.append(i.text)

            for i in par:
                if len(i.text) > 50:
                    name = i.find('b')
                    if name != None:
                        break
            name = name.text

            # =================== Downlading Image =================

            img = soup.find('a', {'class': 'image'})
            url = "https://en.wikipedia.org/" + img.get('href')

            response = requests.get(url).content
            soup = BeautifulSoup(response, 'html.parser')

            img = "https:"+soup.find('a', {'class': 'internal'}).get('href')
            img_data = requests.get(img).content

            path = f'data/{name}.png'
            with open(path, 'wb') as handler:
                handler.write(img_data)

            messagebox.destroy()
            # ===================== Inserting text area and to show the result ================
            text_area = Text(self.frame, height=15,
                             width=70, font=("Calibri 12"))
            text_area.place(x=50, y=320)

            for i in self.paragraph:
                text_area.insert(END, i+"\n")
            text_area['state'] = DISABLED

            # ================ Inserting Image, Name label and save file Button of a person =====================
            img = Image.open(str(path)).resize((250, 250), Image.ANTIALIAS)
            photo = ImageTk.PhotoImage(img)
            label = Label(self.frame, image=photo)
            label.image = photo
            label.place(x=650, y=320)

            name_la = Label(self.frame, text=name, font=(
                "verdana", 11, "bold"), bg="lightblue", borderwidth=5, relief=SUNKEN, width=25, height=2).place(x=650, y=570)

            save_btn = Button(self.frame, text="Save File", font=(
                "verdana 16 bold"), borderwidth=5, relief=SUNKEN, command=lambda: self.save_doc(path), bg="red", fg="white")
            save_btn.place(x=800, y=200, height=50)
        except Exception as e:
            print(e)
            msgbox.showerror("Error", "Person Not Found !! Check the Name")

    # =================== Saving as word file ====================
    def save_doc(self, path):
        doc = docx.Document()
        name = path[5:].replace(".png", "")
        doc.add_heading(f"About {name.capitalize()}", 0)
        doc.add_picture(path)

        for i in self.paragraph:
            doc.add_paragraph(i)

        doc.save(f'{path}.docx')
        msgbox.showinfo(
            "File Saved", f"{name} File Saved Succesfully in data folder!!")


if __name__ == "__main__":
    root = Tk()
    root.geometry("1000x650")
    root.title("Search things That you Want by 21BCA008 Pratham Rathod")
    root.config(bg="Black")
    root.resizable(0, 0)
    obj = wiki(root)
    root.mainloop()